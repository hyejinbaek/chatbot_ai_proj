import os
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.document_loaders import PyMuPDFLoader, PyPDFLoader, TextLoader
from flask import Flask, request, jsonify
import unicodedata
from langchain.schema import Document
from docx import Document as DocxDocument

app = Flask(__name__)

load_dotenv()

openai_api_key = os.getenv("OPENAI_API_KEY")
fine_tuned_model = 'ft:gpt-4o-2024-08-06:auton::ASh95CCN'

def load_docx(file_path):
    # Use the python-docx library to load .docx files
    doc = DocxDocument(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Return a list of langchain Document objects with page_content
    return [Document(page_content="\n".join(full_text))]


def load_document(file_path):
    # 파일 경로의 특수 문자 처리: 정규화
    file_path = unicodedata.normalize('NFKC', file_path)
    
    if file_path.endswith('.pdf'):
        loader = PyMuPDFLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.txt'):
        loader = TextLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.docx'):
        documents = load_docx(file_path)
    else:
        raise ValueError("지원되지 않는 파일 형식입니다.")
    
    if not documents:
        print(f"{file_path} 파일에는 내용이 없습니다.")
    else:
        print(f"{file_path} 파일 내용 로드됨.")
    
    return documents

# 폴더 내 모든 파일을 로드하는 함수 정의
def load_all_documents_in_folder(folder_path):
    docs = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if not filename.endswith(('.pdf', '.txt', '.docx')):
            continue  # PDF, TXT, DOCX 파일만 처리
        try:
            documents = load_document(file_path)
            if documents:
                docs.extend(documents)
                print(f"{filename} 파일이 성공적으로 로드되었습니다.")
            else:
                print(f"{filename} 파일을 로드했지만 내용이 없습니다.")
        except Exception as e:
            print(f"{filename} 파일을 로드하는 중 오류 발생: {e}")
    return docs

# 폴더 경로 지정
folder_path = "./dataset"

# 폴더 내 모든 파일을 로드하여 docs 리스트에 저장
docs = load_all_documents_in_folder(folder_path)
print(f"총 문서 수: {len(docs)}")

# 문서가 없으면 오류를 발생시키지 않도록 체크
if not docs:
    print("로드된 문서가 없습니다. 문제를 확인해 주세요.")
else:
    # 단계 2: 문서 분할(Split Documents)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=50)
    split_documents = text_splitter.split_documents(docs)
    print(f"문서 분할 완료: {len(split_documents)} 개 문서")

    # 단계 3: 임베딩(Embedding) 생성
    embeddings = OpenAIEmbeddings()

    # 단계 4: DB 생성(Create DB) 및 저장
    # 벡터스토어를 생성합니다.
    vectorstore = FAISS.from_documents(documents=split_documents, embedding=embeddings)

    # 단계 5: 검색기(Retriever) 생성
    # 문서에 포함되어 있는 정보를 검색하고 생성합니다.
    retriever = vectorstore.as_retriever()

    # 검색기에 쿼리를 날려 검색된 chunk 결과를 확인합니다.
    retriever.invoke("오토앤 휴가 제도는 무엇인가요?")

    # 단계 6: 프롬프트 생성(Create Prompt)
    prompt = PromptTemplate.from_template(
        """
        이 AI 챗봇은 오토앤 회사의 내부 정책과 절차에 대한 질문에 답변하기 위해 설계되었습니다.
        직원(사용자)들이 회사 규정, 휴가 제도, 복리후생, 또는 기타 회사 관련 사항에 대해 질문하면, 
        주어진 맥락을 기반으로 정확하고 신뢰성 있는 답변을 제공하세요.

        **프롬프트 사용 규칙**:
        1. 사용자가 간단한 인사를 하면, 정중하고 따뜻한 어조로 응답합니다. 예를 들어, 
        - "안녕?" → "안녕하세요! 오늘도 좋은 하루 되세요."
        - "고마워" → "천만에요! 도움이 되었다니 기쁩니다."
        2. 사용자가 오토앤의 특정 정책이나 절차를 물어볼 경우:
        - 질문과 관련된 정보가 컨텍스트에 포함되어 있다면 이를 활용해 명확하고 간결하게 설명하세요.
        - 예를 들어: "오토앤 휴가 제도는 무엇인가요?"라는 질문에 대해 컨텍스트에 정보가 있다면,
            "오토앤의 휴가 제도는 연차, 병가, 특별 휴가로 구성되어 있습니다. 추가적으로 궁금한 사항이 있다면 알려주세요!"와 같이 응답합니다.
        3. 질문과 관련된 정보가 컨텍스트에 없거나 모호할 경우:
        - 친절한 태도로 사용자에게 직접 경영지원본부 또는 적절한 부서로 문의할 것을 안내합니다.
        - 예를 들어: "이 부분에 대한 정확한 정보를 제공하기 어렵습니다. 경영지원본부에 문의해 주시면 도움을 받으실 수 있습니다."
        4. 응답 시 사용자가 이해하기 쉬운 언어를 사용하며, 과도한 전문 용어나 복잡한 설명을 피합니다.
        5. 한국어로만 응답하세요.

        **주의사항**:
        - 응답은 너무 길거나 복잡하지 않게 간결하고 친절하게 작성합니다.
        - 오답을 줄이기 위해 컨텍스트 정보가 명확하지 않으면 반드시 문의를 유도하세요.
        - 모든 응답은 따뜻하고 협조적인 태도로 제공되어야 합니다.

    #Question: 
    {question} 
    #Context: 
    {context} 

    #Answer:"""
    )

    # 단계 7: 언어모델(LLM) 생성
    llm = ChatOpenAI(model_name=fine_tuned_model, temperature=0)

    # 단계 8: 체인(Chain) 생성
    chain = (
        {"context": retriever, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )

    def ask_question(question):
        try:
            # LangChain 모델 호출
            response = chain.invoke(question)
            
            # AI 응답을 대화 기록에 추가 (if needed, depending on your design)
            # question.append(response)  # This line might not be correct if `question` is a string

            return response  # Return the response directly

        except Exception as e:
            print(f"ERROR: 모델 호출 오류 - {e}")
            return "An error occurred while processing your request."


    @app.route('/')
    def index():
        return "API is running. Use /ask to interact with the chatbot.", 200

    @app.route('/ask', methods=['POST'])
    def ask_chatbot():
        data = request.get_json()
        question = data.get("question")

        if not question:
            return jsonify({"error": "No question provided"}), 400

        answer = ask_question(question)
        return jsonify({"answer": answer})

    if __name__ == '__main__':
        app.run(host='0.0.0.0', port=5000)


