# 폴더 전체 읽어와서 langchain
# openai finetuning gpt때는 openai version이 더 낮음
# pip install --upgrade openai

import os
import json
import uuid
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.document_loaders import PyMuPDFLoader, TextLoader
from langchain.schema import Document
from docx import Document as DocxDocument
import unicodedata
from datetime import datetime

load_dotenv()

# OpenAI API 키 가져오기
openai_api_key = os.getenv("OPENAI_API_KEY")

# Fine-tuned 모델 이름
fine_tuned_model = 'ft:gpt-4o-2024-08-06:auton::AUpCFSDU'

# 대화 기록 저장 경로
conversation_log_path = "conversation_logs.json"

# 대화 기록 로드 또는 초기화
if os.path.exists(conversation_log_path):
    with open(conversation_log_path, 'r', encoding='utf-8') as f:
        conversation_logs = json.load(f)
else:
    conversation_logs = {}

# 새로운 세션 ID 생성
def generate_session_id():
    return str(uuid.uuid4())

# 대화 기록 저장 함수
def save_conversation_log(session_id, question, answer):
    if session_id not in conversation_logs:
        conversation_logs[session_id] = []
    conversation_logs[session_id].append({
        "timestamp": datetime.now().isoformat(),
        "question": question,
        "answer": answer
    })
    with open(conversation_log_path, 'w', encoding='utf-8') as f:
        json.dump(conversation_logs, f, ensure_ascii=False, indent=4)

# 이전 대화 기록 불러오기
def get_recent_context(session_id, max_history=5):
    if session_id in conversation_logs:
        return conversation_logs[session_id][-max_history:]
    return []

# 대화 기록 포맷팅
def format_context(recent_context):
    context = ""
    for idx, entry in enumerate(recent_context):
        context += f"[대화 {idx+1}] 질문: {entry['question']}, 답변: {entry['answer']}\n"
    return context

# .docx 파일 로드
def load_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = [paragraph.text for paragraph in doc.paragraphs]
    return [Document(page_content="\n".join(full_text))]

# 문서 로드 함수 (PDF, TXT, DOCX 지원)
def load_document(file_path):
    file_path = unicodedata.normalize('NFKC', file_path)
    if file_path.endswith('.pdf'):
        loader = PyMuPDFLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.txt'):
        loader = TextLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.docx'):
        documents = load_docx(file_path)
    else:
        raise ValueError("지원되지 않는 파일 형식입니다.")
    return documents

# 폴더 내 모든 문서 로드
def load_all_documents_in_folder(folder_path):
    docs = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(('.pdf', '.txt', '.docx')):
            try:
                docs.extend(load_document(file_path))
            except Exception as e:
                print(f"{filename} 로드 중 오류: {e}")
    return docs

# 폴더 내 문서 로드
folder_path = "../dataset"
docs = load_all_documents_in_folder(folder_path)

# 문서 분할
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=50)
split_documents = text_splitter.split_documents(docs)

# 임베딩 및 DB 생성
embeddings = OpenAIEmbeddings()
vectorstore = FAISS.from_documents(documents=split_documents, embedding=embeddings)
retriever = vectorstore.as_retriever()

# 프롬프트 템플릿 생성
prompt = PromptTemplate.from_template(
    """
This is an AI chatbot (assistant) for the company, 오토앤, designed to answer questions. 
Employees will ask questions related to internal company policies at 오토앤. 
Please use the trained model and the following retrieved context to answer these questions. 
If you don't know the answer, politely suggest they contact the HR department. 
Respond in Korean in a courteous and friendly manner.
When the user greets you or asks a question, respond politely and courteously.

#Question: 
{question} 
#Context: 
{context} 

#Answer:"""
)

# 모델 생성 (fine-tuned GPT-4 모델 사용)
llm = ChatOpenAI(model_name=fine_tuned_model, temperature=0)

# 체인 생성 (질문을 처리하고 대답을 반환)
chain = (
    {"context": retriever, "question": RunnablePassthrough()}
    | prompt
    | llm
    | StrOutputParser()
)

# 사용자 질문 처리 함수
def ask_question(session_id, question):
    recent_context = get_recent_context(session_id)
    formatted_context = format_context(recent_context)
    # 프롬프트에 필요한 데이터를 텍스트 형식으로 구성
    full_input = f"""
    # 최근 대화 기록:
    {formatted_context}

    # 질문: 
    {question} 
    """
    # 체인 실행하여 답변 얻기
    response = chain.invoke(full_input)
    save_conversation_log(session_id, question, response)
    print(f"질문: {question}")
    print(f"답변: {response}")

# 실행 예시
if __name__ == '__main__':
    session_id = generate_session_id()
    print(f"새 세션 ID: {session_id}")
    
    # 예시 질문
    ask_question(session_id, "오토앤에서 환경형 성희롱은 어떤 것을 의미하니?")
    ask_question(session_id, "오토앤 휴가 정책에 대해 알려줘.")