# 폴더 전체 읽어와서 langchain

import os
from dotenv import load_dotenv

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.document_loaders import PyMuPDFLoader, PyPDFLoader, TextLoader
import unicodedata
from langchain.schema import Document
from docx import Document as DocxDocument


load_dotenv()

openai_api_key = os.getenv("OPENAI_API_KEY")

fine_tuned_model = 'ft:gpt-4o-2024-08-06:auton::ASh95CCN'

def load_docx(file_path):
    # Use the python-docx library to load .docx files
    doc = DocxDocument(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Return a list of langchain Document objects with page_content
    return [Document(page_content="\n".join(full_text))]


def load_document(file_path):
    # 파일 경로의 특수 문자 처리: 정규화
    file_path = unicodedata.normalize('NFKC', file_path)
    
    if file_path.endswith('.pdf'):
        loader = PyMuPDFLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.txt'):
        loader = TextLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.docx'):
        documents = load_docx(file_path)
    else:
        raise ValueError("지원되지 않는 파일 형식입니다.")
    
    if not documents:
        print(f"{file_path} 파일에는 내용이 없습니다.")
    else:
        print(f"{file_path} 파일 내용 로드됨.")
    
    return documents

# 폴더 내 모든 파일을 로드하는 함수 정의
def load_all_documents_in_folder(folder_path):
    docs = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if not filename.endswith(('.pdf', '.txt', '.docx')):
            continue  # PDF, TXT, DOCX 파일만 처리
        try:
            documents = load_document(file_path)
            if documents:
                docs.extend(documents)
                print(f"{filename} 파일이 성공적으로 로드되었습니다.")
            else:
                print(f"{filename} 파일을 로드했지만 내용이 없습니다.")
        except Exception as e:
            print(f"{filename} 파일을 로드하는 중 오류 발생: {e}")
    return docs


# 폴더 경로 지정
folder_path = "../dataset"

# 폴더 내 모든 파일을 로드하여 docs 리스트에 저장
docs = load_all_documents_in_folder(folder_path)
print(f"총 문서 수: {len(docs)}")

# 단계 2: 문서 분할(Split Documents)
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=50)
split_documents = text_splitter.split_documents(docs)

# 단계 3: 임베딩(Embedding) 생성
embeddings = OpenAIEmbeddings()

# 단계 4: DB 생성(Create DB) 및 저장
# 벡터스토어를 생성합니다.
vectorstore = FAISS.from_documents(documents=split_documents, embedding=embeddings)

# 단계 5: 검색기(Retriever) 생성
# 문서에 포함되어 있는 정보를 검색하고 생성합니다.
retriever = vectorstore.as_retriever()

# 검색기에 쿼리를 날려 검색된 chunk 결과를 확인합니다.
retriever.invoke("오토앤 휴가 제도는 무엇인가요?")

# 단계 6: 프롬프트 생성(Create Prompt)
# 프롬프트를 생성합니다.
prompt = PromptTemplate.from_template(
    """
This is an AI chatbot (assistant) for the company, 오토앤, designed to answer questions. 
Employees will ask questions related to internal company policies at 오토앤. 
Please use the trained model and the following retrieved context to answer these questions. 
If you don't know the answer, politely suggest they contact the HR department. 
Respond in Korean in a courteous and friendly manner.
When the user greets you or asks a question, respond politely and courteously.

#Question: 
{question} 
#Context: 
{context} 

#Answer:"""
)

# 단계 7: 언어모델(LLM) 생성
# 모델(LLM) 을 생성합니다.
llm = ChatOpenAI(model_name=fine_tuned_model, temperature=0)

# 단계 8: 체인(Chain) 생성
chain = (
    {"context": retriever, "question": RunnablePassthrough()}
    | prompt
    | llm
    | StrOutputParser()
)

# 사용자 입력 받기 및 체인 실행
def ask_question(question):
    response = chain.invoke(question)
    print(f"질문: {question}")
    print(f"답변: {response}")

if __name__ == '__main__':
    # 예시 질문
    ask_question("오토앤에서 환경형 성희롱은 어떤 것을 의미하니?")