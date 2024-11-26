# gpt-4o-mini + langchain

import os
import json
import uuid
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.document_loaders import PyMuPDFLoader, PyPDFLoader, TextLoader
from flask import Flask, request, jsonify, session
import unicodedata
from langchain.schema import Document
from docx import Document as DocxDocument
from datetime import datetime
import pandas as pd

app = Flask(__name__)

load_dotenv()
app.secret_key = os.getenv('FLASK_SECRET_KEY')
openai_api_key = os.getenv("OPENAI_API_KEY")
fine_tuned_model = 'gpt-4o-mini'

# 대화 기록 저장 경로
def get_conversation_log_path():
    today_date = datetime.now().strftime("%Y%m%d")
    return f"./log/conversation_logs_{today_date}.json"

# 대화 기록 로드 또는 초기화
def load_conversation_logs():
    conversation_log_path = get_conversation_log_path()
    if os.path.exists(conversation_log_path):
        with open(conversation_log_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

conversation_logs = load_conversation_logs()

# 새로운 세션 ID 생성
def generate_session_id():
    return str(uuid.uuid4())

# 대화 기록 저장 함수
def save_conversation_log(session_id, question, answer):
    if session_id not in conversation_logs:
        conversation_logs[session_id] = []
    conversation_logs[session_id].append({
        "timestamp": datetime.now().isoformat(),
        "question": question,
        "answer": answer
    })
    
    # 날짜별 대화 기록 파일에 저장
    conversation_log_path = get_conversation_log_path()
    with open(conversation_log_path, 'w', encoding='utf-8') as f:
        json.dump(conversation_logs, f, ensure_ascii=False, indent=4)

# 이전 대화 기록 불러오기
def get_recent_context(session_id, max_history=5):
    if session_id in conversation_logs:
        return conversation_logs[session_id][-max_history:]
    return []

def format_context(recent_context):
    context = ""
    for idx, entry in enumerate(recent_context):
        context += f"[대화 {idx+1}] 질문: {entry['question']}, 답변: {entry['answer']}\n"
    return context

# 엑셀 파일 로드 함수
def load_excel(file_path):
    print(f"Loading Excel file: {file_path}")  # 디버깅: 엑셀 파일 경로 출력
    try:
        # pandas로 엑셀 파일을 로드하여 데이터프레임으로 변환
        df = pd.read_excel(file_path, engine='openpyxl')  # 엑셀 파일 읽기
        print(f"Excel file loaded successfully: {file_path}")  # 디버깅: 로드 성공 출력
        # 각 행을 문서로 변환하여 페이지 콘텐츠에 추가
        full_text = []
        for index, row in df.iterrows():
            full_text.append(" ".join(str(cell) for cell in row))  # 각 셀의 값을 공백으로 구분하여 결합
        return [Document(page_content="\n".join(full_text))]
    except Exception as e:
        print(f"Error loading Excel file: {e}")  # 디버깅: 오류 출력
        return []

# .docx 파일 로드
def load_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return [Document(page_content="\n".join(full_text))]

def load_document(file_path):
    file_path = unicodedata.normalize('NFKC', file_path)
    print(f"Checking file type: {file_path}")  # 디버깅: 파일 유형 확인
    if file_path.endswith('.pdf'):
        loader = PyMuPDFLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.txt'):
        loader = TextLoader(file_path=file_path)
        documents = loader.load()
    elif file_path.endswith('.docx'):
        documents = load_docx(file_path)
    elif file_path.endswith('.xlsx'):
        documents = load_excel(file_path)
    else:
        raise ValueError("지원되지 않는 파일 형식입니다.")
    return documents

loaded_files = {}  # 파일 수정 시간 기록을 위한 딕셔너리


def load_all_documents_in_folder(folder_path):
    docs = []
    print(f"Listing files in folder: {folder_path}")  # 디버깅: 폴더 내 파일 목록 출력
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(('.pdf', '.txt', '.docx', '.xlsx')):  # 엑셀 파일도 포함
            try:
                # 새 파일이나 수정된 파일만 로드
                print(f"Loading file: {file_path}")  # 디버깅: 파일 경로 출력
                if file_path not in loaded_files or os.path.getmtime(file_path) > loaded_files[file_path]:
                    docs.extend(load_document(file_path))
                    loaded_files[file_path] = os.path.getmtime(file_path)
            except Exception as e:
                print(f"{filename} 로드 중 오류: {e}")
    return docs

# 문서 로드 및 처리
folder_path = "./dataset"
docs = load_all_documents_in_folder(folder_path)
print(f"총 문서 수: {len(docs)}")

for doc in docs:
    #print(doc.page_content)  # 문서 내용 출력
    # 문서 내용을 파일로 저장하는 코드
    with open("document_contents.txt", "w", encoding="utf-8") as file:
        for doc in docs:
            # 문서의 내용을 파일에 기록
            file.write(doc.page_content + "\n\n")  # 문서 내용 뒤에 줄바꿈 추가

if not docs:
    print("로드된 문서가 없습니다. 문제를 확인해 주세요.")
else:
    # 문서를 청크로 나누기 위한 분할기 설정
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,  # 청크 크기 (조정 가능)
        chunk_overlap=200,  # 청크가 겹치는 부분 (조정 가능)
        length_function=len,
    )
    split_documents = text_splitter.split_documents(docs)
    print(f"문서 분할 완료: {len(split_documents)} 개 문서")
    

    embeddings = OpenAIEmbeddings()

    # 벡터 스토어 경로 설정
    vectorstore_path = "./vector/faiss_index"

    # 벡터 스토어 로드 또는 생성
    if os.path.exists(vectorstore_path):
        print("벡터 스토어 로드 중...")
        vectorstore = FAISS.load_local(
            vectorstore_path, 
            embeddings=embeddings, 
            allow_dangerous_deserialization=True  # 안전한 파일이라 확신할 경우 설정
        )
    else:
        print("벡터 스토어 생성 중...")
        vectorstore = FAISS.from_documents(documents=split_documents, embedding=embeddings)
        vectorstore.save_local(vectorstore_path)

    retriever = vectorstore.as_retriever(k=3)

    results = retriever.invoke("오토앤 직급체계는 무엇인가요?")
    print("검색 결과:", results)
    

    prompt = PromptTemplate.from_template(
        """
        이 AI 챗봇은 오토앤 회사의 내부 정책과 절차에 대한 질문에 답변하기 위해 설계되었습니다.
        직원(사용자)들이 회사 규정, 휴가 제도(연차), 복리후생, 또는 기타 회사 관련 사항에 대해 질문하면, 
        주어진 맥락을 기반으로 정확하고 신뢰성 있는 답변을 제공하세요.

        **프롬프트 사용 규칙**:
        1. 사용자가 간단한 인사를 하면, 정중하고 따뜻한 어조로 응답합니다. 예를 들어, 
        - "안녕?" → "안녕하세요! 오늘도 좋은 하루 되세요."
        - "고마워" → "천만에요! 도움이 되었다니 기쁩니다." 외에도 다양한 표현으로 따뜻하고 친절하게 응답하세요.
        2. 사용자가 오토앤의 특정 정책이나 절차를 물어볼 경우:
        - 질문과 관련된 정보가 컨텍스트에 포함되어 있다면 이를 활용해 명확하고 간결하게 설명하세요.
        - 예를 들어: "오토앤 휴가 제도는 무엇인가요?"라는 질문에 대해 컨텍스트에 정보가 있다면,
            "오토앤의 휴가 제도는 연차, 병가, 특별 휴가로 구성되어 있습니다. 추가적으로 궁금한 사항이 있다면 알려주세요!"와 같이 응답합니다.
        - 사용자의 질문이 명확하지 않거나 키워드가 부족한 경우, 구체적으로 질문할 수 있도록 친절히 안내하세요.
        - 예를 들어 : "오토앤에서 일할 때 규정이 있나요?"라고 물었다면,
            "오토앤에서 적용되는 규정을 말씀하시는 건가요? 예를 들어, 근무 시간, 사내 동호회, 혹은 연차와 관련된 규정을 구체적으로 질문해 주시면 더 정확한 답변을 드릴 수 있습니다."
        - 사용자가 질문을 더 명확히 할 수 있도록 다음과 같이 추가적인 예시를 제시하세요 : 
            "근무 시간 관련 규정인가요?", "사내 동호회에 대한 정보가 필요하신가요?"
        3. 질문과 관련된 정보가 컨텍스트에 없거나 모호할 경우:
        - 친절한 태도로 사용자에게 직접 경영지원본부 또는 적절한 부서로 문의할 것을 안내합니다.
        - 예를 들어: "이 부분에 대한 정확한 정보를 제공하기 어렵습니다. 경영지원본부 또는 관련 부서에 문의해 주시면 도움을 받으실 수 있습니다."
        4. 응답 시 사용자가 이해하기 쉬운 언어를 사용하며, 과도한 전문 용어나 복잡한 설명을 피합니다.
        5. 한국어로만 응답하세요.
        6. 사용자가 너의(AI챗봇) 존재를 묻는다면 친절하게 설명해주세요. 예를 들어,
        - 오토앤에 대해 궁금한 점이 있으시군요! 회사의 정책, 절차, 복리후생, 휴가 제도 등 다양한 정보를 제공할 수 있습니다. 어떤 특정한 정보가 필요하신가요? 예를 들어, "휴가 제도에 대해 알고 싶어요" 또는 "복리후생이 어떤 게 있나요?"와 같이 구체적으로 질문해 주시면 더 정확한 답변을 드릴 수 있습니다.
        와 같이 친절하게 다양한 표현으로 설명하세요.
        
        **주의사항**:
        - 응답은 너무 길거나 복잡하지 않게 간결하고 친절하게 작성합니다.
        - 오답을 줄이기 위해 컨텍스트 정보가 명확하지 않으면 반드시 문의를 유도하세요.
        - 모든 응답은 따뜻하고 협조적인 태도로 제공되어야 합니다.

        #Question: 
        {question} 
        #Context: 
        {context} 

        #Answer:""" 
    )

    llm = ChatOpenAI(model_name=fine_tuned_model, temperature=0)

    chain = (
        {"context": retriever, "question": RunnablePassthrough()}
        | prompt
        | llm
        | StrOutputParser()
    )

    def ask_question(session_id, question):
        recent_context = get_recent_context(session_id)
        formatted_context = format_context(recent_context)
        # print(f"Searching for context with question: {question}")
        print(f"Context found: {formatted_context}")

        full_input = f"""
        # 최근 대화 기록:
        {formatted_context}

        # 질문: 
        {question} 
        """
        try:
            response = chain.invoke(full_input)
            print("프롬프트 입력값:", full_input)
            save_conversation_log(session_id, question, response)
            return response
        except Exception as e:
            print(f"ERROR: 모델 호출 오류 - {e}")
            return "An error occurred while processing your request."

    @app.route('/')
    def index():
        return "API is running. Use /ask to interact with the chatbot.", 200

    
    @app.route('/ask', methods=['POST'])
    def ask_chatbot():
        data = request.get_json()
        question = data.get("question")
        
        # 디버깅: 세션 ID 확인
        if 'session_id' not in session:
            print("새로운 세션 ID 생성")
            session['session_id'] = str(uuid.uuid4())  # 고유한 세션 ID 생성

        session_id = session['session_id']
        print(f"Session ID: {session_id}")  # 디버깅 로그
        
        if not question:
            return jsonify({"error": "No question provided"}), 400

        answer = ask_question(session_id, question)

        return jsonify({
            "answer": answer,
            "session_id": session_id
        })
    
    if __name__ == '__main__':
        app.run(host='0.0.0.0', port=5000)
